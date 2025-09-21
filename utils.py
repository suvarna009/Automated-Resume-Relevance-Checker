# # parser_utils.py
# import fitz  # PyMuPDF
# import docx
# import re

# def extract_text_from_pdf_bytes(file_bytes):
#     """
#     file_bytes: bytes-like (uploaded_file.read())
#     returns: string text
#     """
#     try:
#         doc = fitz.open(stream=file_bytes, filetype="pdf")
#         text = []
#         for page in doc:
#             text.append(page.get_text())
#         return "\n".join(text)
#     except Exception:
#         return ""

# def extract_text_from_docx_bytes(file_bytes):
#     """
#     file_bytes: bytes-like (uploaded_file.read())
#     docx.Document can accept a file-like object if wrapped
#     """
#     try:
#         # docx expects a path or a file-like object. We'll decode bytes via BytesIO
#         from io import BytesIO
#         f = BytesIO(file_bytes)
#         doc = docx.Document(f)
#         paras = [p.text for p in doc.paragraphs]
#         return "\n".join(paras)
#     except Exception:
#         return ""

# def clean_text(text):
#     if not text:
#         return ""
#     # Normalize whitespace and lower
#     text = re.sub(r'\s+', ' ', text).strip()
#     return text

# def parse_uploaded_file(uploaded_file):
#     """
#     uploaded_file: streamlit uploaded file
#     return: parsed plain text (lowercased, cleaned)
#     """
#     if uploaded_file is None:
#         return ""
#     name = uploaded_file.name.lower()
#     raw = uploaded_file.read()
#     if name.endswith(".pdf"):
#         txt = extract_text_from_pdf_bytes(raw)
#     elif name.endswith(".docx") or name.endswith(".doc"):
#         txt = extract_text_from_docx_bytes(raw)
#     else:
#         try:
#             txt = raw.decode("utf-8", errors="ignore")
#         except Exception:
#             txt = ""
#     return clean_text(txt).lower()
# utils.py
import re
from io import BytesIO
from collections import Counter
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from PyPDF2 import PdfReader
from docx import Document
import math

# ---------- file parsing ----------
def extract_text_from_pdf_bytes(b):
    text_chunks = []
    try:
        reader = PdfReader(BytesIO(b))
        for p in reader.pages:
            page_text = p.extract_text() or ""
            text_chunks.append(page_text)
    except Exception:
        # fallback: try decode
        try:
            text_chunks.append(b.decode(errors="ignore"))
        except:
            pass
    return "\n".join(text_chunks)

def extract_text_from_docx_bytes(b):
    try:
        doc = Document(BytesIO(b))
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception:
        try:
            return b.decode(errors="ignore")
        except:
            return ""

def parse_uploaded_file(uploaded_file):
    """
    streamed_file: streamlit uploaded_file object
    returns plain text extracted
    """
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    b = uploaded_file.read()
    if name.endswith(".pdf"):
        return extract_text_from_pdf_bytes(b)
    elif name.endswith(".docx") or name.endswith(".doc"):
        return extract_text_from_docx_bytes(b)
    elif name.endswith(".txt"):
        try:
            return b.decode(errors="ignore")
        except:
            return ""
    else:
        # best-effort
        try:
            return extract_text_from_pdf_bytes(b)
        except:
            try:
                return extract_text_from_docx_bytes(b)
            except:
                try:
                    return b.decode(errors="ignore")
                except:
                    return ""

# ---------- keyword extraction ----------
def extract_keywords(text, top_n=12):
    text = (text or "").strip()
    if not text:
        return []
    vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1,2), max_features=200)
    try:
        tfidf = vectorizer.fit_transform([text])
        feature_names = vectorizer.get_feature_names_out()
        scores = tfidf.toarray()[0]
        top_indices = scores.argsort()[::-1][:top_n]
        keywords = [feature_names[i] for i in top_indices if scores[i] > 0]
        return keywords
    except Exception:
        # fallback: simple token frequency
        tokens = re.findall(r"[A-Za-z0-9\+\#\.\-]+", text.lower())
        commons = [w for w,c in Counter(tokens).most_common(top_n)]
        return commons

# ---------- matching & feedback ----------
def compute_match_and_feedback(resume_text, jd_text):
    resume_text = (resume_text or "")
    jd_text = (jd_text or "")
    # similarity (tfidf cosine)
    try:
        vectorizer = TfidfVectorizer(stop_words='english')
        X = vectorizer.fit_transform([jd_text, resume_text])
        sim = float(cosine_similarity(X[0:1], X[1:2])[0][0])
    except Exception:
        sim = 0.0

    jd_keywords = extract_keywords(jd_text, top_n=10)
    matched = []
    missing = []
    resume_lower = resume_text.lower()
    for kw in jd_keywords:
        if kw.lower() in resume_lower:
            matched.append(kw)
        else:
            missing.append(kw)

    skill_match_pct = len(matched) / max(1, len(jd_keywords))
    # weighted final score: 70% overall text similarity + 30% skill presence
    final_score = (0.7 * sim) + (0.3 * skill_match_pct)
    final_score_percent = round(final_score * 100, 1)

    # Build feedback (human-friendly)
    feedback_lines = []
    if final_score_percent >= 85:
        feedback_lines.append("🎉 Excellent match — your resume aligns very well with the JD.")
    elif final_score_percent >= 65:
        feedback_lines.append("🙂 Good match, but you can improve. Highlight measurable achievements and missing skills below.")
    elif final_score_percent >= 45:
        feedback_lines.append("⚠️ Fair match — focus on adding the missing skills and tailor the summary/skills section.")
    else:
        feedback_lines.append("🔴 Low match — you should tailor your resume to this job and add the missing skills and achievements.")

    if missing:
        feedback_lines.append("🔎 **Missing / weakly represented skills**: " + ", ".join(missing))
        feedback_lines.append("✅ **How to fix**: add these skills to a prominent 'Skills' section and show 1–2 bullet points per skill demonstrating where you used it (metrics help).")
    else:
        feedback_lines.append("✅ All top JD keywords are present in the resume — great!")

    # Formatting suggestions
    word_count = len(re.findall(r"\w+", resume_text))
    if word_count < 200:
        feedback_lines.append("📝 Resume may be too short — add more concrete bullets and metrics.")
    elif word_count > 1400:
        feedback_lines.append("📝 Resume may be too long — trim irrelevant experience, aim for clarity.")
    else:
        feedback_lines.append("📝 Resume length looks fine.")

    # Sample micro-rewrite suggestions (examples)
    if missing:
        example_suggestions = []
        for s in missing[:5]:
            example_suggestions.append(f"- Add a bullet: 'Implemented {s} to achieve X% improvement in Y' (quantify if possible).")
        feedback_lines.append("✍️ Example bullets you could add:\n" + "\n".join(example_suggestions))

    return {
        "score": final_score_percent,
        "similarity": round(sim * 100, 1),
        "skill_match_pct": round(skill_match_pct * 100, 1),
        "jd_keywords": jd_keywords,
        "matched": matched,
        "missing": missing,
        "feedback_lines": feedback_lines,
        "word_count": word_count
    }
